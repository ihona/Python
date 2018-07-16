import os
import docx
# pip install python-docx
#用于解析word文件
file='D:\word.html'
if os.path.exists(file):
    os.remove(file)

def cat():
    doc = docx.Document('D:\word.docx')
    word = ""
    for num in range(2,len(doc.paragraphs)):
        numtext = doc.paragraphs[num].text
        nn = numtext[-30:]
        # nn = numtext.split("    ")[1].split(r"\n")[0]
        word +="<tr><td>"+ numtext.split("    ")[0] + "</td><td>" + nn + "</td></tr>"
        num += 1

        # print(str.split("[")[1].split("]")[0])

    f = open(file, "a", encoding='utf-8')
    data="<table>"+word+"</table>"
    f.write(data)
    f.close()

cat()

